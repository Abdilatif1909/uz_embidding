from docx import Document
from docx.shared import Pt
import re

INPUT = "MAQOLA.md"
OUTPUT = "MAQOLA.docx"

def add_code_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)


def main():
    doc = Document()
    in_code = False
    with open(INPUT, 'r', encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')
            if line.strip().startswith('```'):
                in_code = not in_code
                if in_code:
                    code_acc = []
                else:
                    add_code_paragraph(doc, '\n'.join(code_acc))
                continue

            if in_code:
                code_acc.append(line)
                continue

            # Headings
            m = re.match(r'^(#{1,6})\s*(.*)$', line)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                # docx heading levels go from 0..9 (we use 1..5)
                doc.add_heading(text, level=min(level, 5))
                continue

            # Blockquote
            if line.strip().startswith('>'):
                text = line.strip().lstrip('> ').strip()
                p = doc.add_paragraph(text)
                p.runs[0].italic = True
                continue

            # Lists (unordered)
            if re.match(r'^\s*[-\*]\s+.+$', line):
                item = re.sub(r'^\s*[-\*]\s+', '', line)
                doc.add_paragraph(item, style='List Bullet')
                continue

            # Numbered lists
            if re.match(r'^\s*\d+\.\s+.+$', line):
                item = re.sub(r'^\s*\d+\.\s+', '', line)
                doc.add_paragraph(item, style='List Number')
                continue

            # Horizontal rule
            if re.match(r'^[\-\*_]{3,}\s*$', line):
                doc.add_page_break()
                continue

            # Empty line -> paragraph break
            if line.strip() == '':
                doc.add_paragraph('')
                continue

            # Regular paragraph
            doc.add_paragraph(line)

    doc.save(OUTPUT)
    print(f"✅ Converted: {INPUT} -> {OUTPUT}")

if __name__ == '__main__':
    main()
